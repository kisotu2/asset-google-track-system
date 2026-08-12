from copy import deepcopy
from docx import Document

SOURCE = "/Users/kiso2/Library/CloudStorage/OneDrive-StrathmoreUniversity/Lemayian_152170.docx"
OUTPUT = "Smart_Asset_Management_with_Google_Device_Tracking.docx"

doc = Document(SOURCE)

def paragraph_starts(prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")

def replace(prefix, text):
    paragraph_starts(prefix).text = text

def insert_after(anchor, text, style=None):
    new = anchor._parent.add_paragraph(text, style=style or anchor.style)
    anchor._p.addnext(new._p)
    return new

def insert_before(anchor, text, style=None):
    new = anchor._parent.add_paragraph(text, style=style or anchor.style)
    anchor._p.addprevious(new._p)
    return new

# Title and abstract
replace("Machine Learning Enabled Smart Asset Management System", "Machine Learning Enabled Smart Asset Management System for Predictive Maintenance, Asset Lifecycle Management and Google-Based Device Location Tracking")
replace("The Smart Asset Management System is a web-based solution", "The Smart Asset Management System is a web-based solution designed to improve the management of ICT assets such as laptops, desktop computers, printers, and other technological equipment within an organisation. Many institutions still rely on manual methods, including spreadsheets and paper records, which make asset assignment, location verification and lifecycle management difficult. The proposed system provides a centralised platform for registering, monitoring and managing assets throughout their lifecycle.")
replace("The system will be developed using PHP", "The system will be developed using PHP, MySQL, HTML, CSS, JavaScript and XAMPP. It will support asset registration, assignment tracking, user and role management, status monitoring, maintenance records and report generation. A Google Maps-based location dashboard will display the last authorised location of a tracked device, its timestamp, accuracy and tracking status. Location records will be submitted only after explicit device-user permission, then stored securely for authorised administrators.")
replace("To enhance security, the system will implement", "To enhance security, the system will implement Two-Factor Authentication (2FA) through email verification using PHPMailer and role-based access control. Machine learning features will support predictive maintenance, smart asset allocation, asset lifecycle prediction, anomaly detection and automated alerts. By analysing device age, repair history, warranty status, usage patterns, assignment records and permitted location events, the system will help administrators identify high-risk assets and make better decisions before problems become serious.")
replace("Overall, the Smart Asset Management System aims", "Overall, the Smart Asset Management System aims to improve accountability, reduce asset loss, support timely maintenance, strengthen security and simplify the day-to-day management of ICT resources. The tracking module is designed as an accountable, consent-based enhancement rather than a covert surveillance tool.")
replace("Keywords: Smart Asset Management System", "Keywords: Smart Asset Management System, ICT Asset Management, Device Location Tracking, Google Maps Platform, Asset Tracking, Predictive Maintenance, Machine Learning, Asset Lifecycle Management, Two-Factor Authentication (2FA), Anomaly Detection, Asset Allocation, Report Generation, PHP, MySQL.")

# Abbreviations
insert_after(paragraph_starts("GLPI –"), "GPS – Global Positioning System")
insert_after(paragraph_starts("HTML –"), "HTTPS – Hypertext Transfer Protocol Secure")

# Chapter 1
bg = paragraph_starts("This project goes beyond basic asset tracking")
insert_after(bg, "The proposed system also introduces authorised device-location tracking. Google Maps Platform will be used to visualise a device's last known coordinates on an administrator dashboard. A device can provide its location through a browser check-in using the Geolocation API, or through a managed mobile client where available. The system will record latitude, longitude, accuracy, capture time and source only after permission has been granted; it will not claim to use Google to remotely locate arbitrary devices.")
replace("This project, therefore, seeks to develop a Smart Asset Management", "This project, therefore, seeks to develop a Smart Asset Management System that automates asset tracking and management while integrating consent-based device-location tracking through Google Maps Platform, machine learning for predictive maintenance, smart asset allocation, anomaly detection, lifecycle prediction, and automated alerts.")
insert_after(paragraph_starts("To integrate security features"), "To implement consent-based device location check-ins, Google Maps visualisation, location history and authorised location alerts")
insert_after(paragraph_starts("How can machine learning be used"), "How can authorised location data and Google Maps visualisation improve device accountability without compromising privacy?")
insert_after(paragraph_starts("This study is justified because it proposes"), "The Google-based tracking capability is justified because an assignment record alone does not always confirm whether a device is at its expected workplace or has checked in recently. A last-known-location view and location history can improve recovery, inventory verification and accountability. The feature will be limited to organisation-owned devices and governed by explicit notice, consent, role-based access, audit logs and a defined retention period.")
replace("This study focuses on the design and development", "This study focuses on the design and development of a web-based Smart Asset Management System for managing ICT assets within an organisation, including laptops, desktops, printers and related computing devices. The system will provide a centralised platform for asset registration, assignment, monitoring, reporting, user management and secure authentication. It will incorporate predictive maintenance, asset lifecycle prediction, anomaly detection, smart asset allocation and a Google Maps-based device-location module. The tracking module will capture an authorised device's last known location through an HTTPS browser check-in or a managed mobile client, display the point and timestamp on a map, retain location history, and demonstrate a configurable approved-area alert. The system will manage asset details, condition, maintenance history, warranty information, assignment records, user activities and permitted location records, using sample data where real organisational data is unavailable. The project excludes covert tracking, continuous background laptop tracking without a device agent, access to Google Find My Device, procurement management, asset depreciation, barcode hardware integration, enterprise resource planning functions and advanced hardware repair diagnostics.")
insert_after(paragraph_starts("Limited machine learning accuracy"), "Location availability and accuracy – Location updates depend on device connectivity, browser or operating-system permissions, GPS/Wi-Fi capability and whether the authorised check-in client is active. The displayed location is therefore a last known location, not a guaranteed real-time position.")
insert_after(paragraph_starts("The machine learning component will focus"), "The location-tracking demonstration will use consent-based location check-ins and last-known-location reporting; it will not implement covert surveillance, remote activation of device location, or integration with consumer Google Find My Device services.")

# Literature review / framework
framework = paragraph_starts("Conceptual Framework")
h = insert_before(framework, "Google Maps Platform and Authorised Device Location Tracking", style="Heading 2")
insert_after(h, "Google Maps Platform provides map visualisation and geospatial services that can be integrated into a web application through an API key. In this project, it will be used to render authorised device-location records as markers and to support the display of approved-area boundaries. The map service is a presentation and geospatial layer; it does not independently discover the location of a laptop or phone. The application must receive coordinates from a permitted device-side source before a marker can be shown.")
insert_after(paragraph_starts("Google Maps Platform provides map visualisation"), "For browser-based check-ins, the web application will use the browser Geolocation API over HTTPS. The browser requests the device user's permission before providing coordinates, and the user may refuse or withdraw that permission. This makes the approach suitable for an academic demonstration of accountable asset verification. A managed Android client could later submit periodic authorised updates using Google Play services, but that capability is outside the core browser-only demonstration. Location data will be protected through least-privilege access, encrypted transport, audit logging, retention limits and a visible privacy notice, consistent with the treatment of location data as personal data under the Kenya Data Protection Act, 2019.")
replace("The conceptual framework for this study explains", "The conceptual framework for this study explains how the proposed Smart Asset Management System will work. The framework is based on six main parts: inputs, core application processes, authorised location capture and Google Maps services, machine learning processing, API integration and outputs, as shown in Figure 1: Conceptual Framework of Smart Asset Management System.")
replace("The main inputs will include asset details", "The main inputs will include asset details, user details, department information, assignment records, asset status, maintenance history, warranty information, system activity logs and permitted location records. A location record contains the asset identifier, latitude, longitude, accuracy, capture time, source and consent status. These inputs provide the data needed for normal asset management, authorised location verification and machine learning analysis.")
replace("The PHP and MySQL-based web application will handle", "The PHP and MySQL-based web application will handle user login, Two-Factor Authentication, role-based access control, asset registration, asset assignment, status updates, report generation, location-history access and alert display. Only authorised roles will be allowed to view locations, and every location view or export will be recorded in the audit log.")
location_heading = insert_after(paragraph_starts("The PHP and MySQL-based web application will handle"), "Authorised Location Capture and Google Maps Services", style="Heading 3")
insert_after(location_heading, "A device user will initiate an HTTPS location check-in and explicitly grant location permission. The web application will validate and store the submitted coordinates and render the most recent valid point on a Google Maps dashboard. Administrators can view a last-known-location marker, capture time, accuracy, location history and an approved-area alert. The system will never request or collect location silently.")
replace("The Flask API will act as the communication bridge", "The Flask API will act as the communication bridge between the PHP web application and the Python machine learning models. When the system needs a prediction, PHP will send relevant asset data to the Flask API. Location records remain in the main application database; only aggregated or minimised location-derived features may be sent to the analytics service where needed for an approved anomaly rule.")
replace("The expected outputs will include accurate asset records", "The expected outputs will include accurate asset records, assignment reports, authorised last-known locations, location-history reports, approved-area alerts, maintenance alerts, lifecycle predictions, anomaly notifications and decision-support reports. These outputs will help administrators manage assets more efficiently, improve accountability, reduce downtime and make better maintenance and replacement decisions.")
replace("This chapter has reviewed the current state", "This chapter has reviewed the current state of ICT asset management, common challenges, existing solutions, machine learning applications, device-location tracking and the gaps that remain. The review shows that organisations need more than a basic asset register. They need systems that are secure, centralised, intelligent and capable of supporting proactive decisions. The proposed system responds by combining a PHP and MySQL asset-management platform, consent-based location capture, Google Maps visualisation and Python-based machine-learning models connected through a Flask API.")

# Methodology and deliverables
replace("Data will be collected from asset records", "Data will be collected from asset records, user assignment records, maintenance history, warranty information, asset status updates, login or system activity logs and authorised location check-ins. Location samples will contain only the fields necessary for tracking: asset ID, coordinates, accuracy, capture time, source and consent status. Where real organisational data is unavailable, realistic sample data will be prepared for development and testing purposes.")
replace("The collected data will be cleaned", "The collected data will be cleaned, validated and organised before being used by the system. This will include removing duplicate records, handling missing values, standardising fields such as asset status and department names, validating coordinate ranges and timestamps, and preparing useful features such as device age, number of repairs, warranty status, assignment duration and time since the last authorised location check-in.")
replace("This project adopts the Object-Oriented Analysis", "This project adopts the Object-Oriented Analysis and Design (OOAD) paradigm to guide the analysis, design and implementation of the Smart Asset Management System. OOAD is appropriate because the system consists of interconnected modules: Authentication and Security, User Management, Asset Registration, Asset Assignment and Tracking, Device Location Tracking, Google Maps Dashboard, Reporting and Analytics, Alert Management and Machine Learning-Based Predictive Analytics. Each module can be modelled as an independent object with defined attributes, methods and responsibilities, promoting modularity, maintainability and scalability.")
insert_after(paragraph_starts("User authentication and authorization"), "Location-record database design and consent/audit controls")
insert_after(paragraph_starts("Asset assignment and tracking"), "Authorised device location check-in and Google Maps dashboard")
insert_after(paragraph_starts("Alert and notification functionality"), "Location-history reports and approved-area alert demonstration")
replace("Sprint planning is conducted at the beginning", "Sprint planning is conducted at the beginning of each development cycle. During this phase, the developer selects backlog items to be completed within the sprint and estimates the effort required. Early sprints focus on database development, authentication and core asset management functionality. Subsequent sprints implement the location-record schema, consent flow, HTTPS browser check-in, Google Maps dashboard, history and alert views. Later sprints focus on machine learning integration, reporting, privacy and security testing, and system evaluation.")
replace("The sprint backlog contains the specific tasks", "The sprint backlog contains the specific tasks selected for implementation during a particular sprint. Examples include creating database tables, implementing user authentication, developing asset-tracking interfaces, implementing the consent and location check-in endpoint, configuring the Google Maps dashboard, training predictive-maintenance models or configuring email-based alerts.")
replace("The integrated increment represents the working version", "The integrated increment represents the working version of the Smart Asset Management System produced at the end of each sprint. The final integrated increment combines user authentication, role management, asset registration, assignment tracking, authorised device-location tracking, Google Maps visualisation, location-history reporting, privacy controls, alert generation, predictive maintenance, asset lifecycle prediction, anomaly detection and dashboard visualisation into a complete intelligent asset management platform.")
insert_after(paragraph_starts("This module will support asset registration"), "Device Location Tracking and Google Maps Module", style="Heading 2")
insert_after(paragraph_starts("Device Location Tracking and Google Maps Module"), "This module will obtain an authorised location check-in, store the asset-linked location record and display the last known device location, timestamp and accuracy on a Google Maps dashboard. It will provide authorised location history, approved-area alerts, tracking enable/disable status and an audit trail of location access.")
insert_after(paragraph_starts("PHPMailer will be used"), "Google Maps Platform", style="Heading 2")
insert_after(paragraph_starts("Google Maps Platform"), "Google Maps JavaScript API will be used to display authorised device locations and approved-area boundaries on the administrator dashboard. The API key will be restricted to the deployed web origin and stored outside source control. The browser Geolocation API will collect a location only after the user grants permission over HTTPS.")

# References inserted before existing references as a grouped addition.
refs = paragraph_starts("References")
insert_after(refs, "Google. (n.d.). Maps JavaScript API usage and billing. Google for Developers. https://developers.google.com/maps/documentation/javascript/usage-and-billing")
insert_after(paragraph_starts("Google. (n.d.). Maps JavaScript API usage"), "Google. (n.d.). FusedLocationProviderClient. Google for Developers. https://developers.google.com/android/reference/com/google/android/gms/location/FusedLocationProviderClient")
insert_after(paragraph_starts("Google. (n.d.). FusedLocationProviderClient"), "Kenya Law. (2019). Data Protection Act, 2019. https://new.kenyalaw.org/akn/ke/act/2019/24/eng@2019-11-15")
insert_after(paragraph_starts("Kenya Law. (2019). Data Protection Act"), "Mozilla Developer Network. (n.d.). Geolocation API. https://developer.mozilla.org/en-US/docs/Web/API/Geolocation_API")

# Add an implementation-ready appendix ahead of the original appendix section.
appendices = paragraph_starts("Appendices")
app_heading = insert_after(appendices, "Appendix 3: Google-Based Device Tracking Specification", style="Heading 1")
insert_after(app_heading, "Purpose: to verify the last authorised location of organisation-owned ICT devices and display it to authorised administrators. The feature is not intended to covertly monitor users or to provide a guaranteed real-time recovery service.")
table_anchor = insert_after(app_heading, "Proposed Location Record Fields", style="Heading 2")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for cell, value in zip(table.rows[0].cells, ["Field", "Description", "Control"]):
    cell.text = value
for values in [
    ("asset_id", "Links the location record to the registered asset.", "Foreign key; required."),
    ("latitude / longitude", "Coordinates received from the permitted device check-in.", "Validate range; encrypt in transit."),
    ("accuracy_meters", "Reported estimate of positional accuracy.", "Show to administrator; reject invalid values."),
    ("captured_at", "Time at which the device captured the location.", "Display as last known, not live location."),
    ("source", "Browser check-in or managed mobile client.", "Record for auditability."),
    ("consent_status", "Permission/notice status at capture.", "Required before storing a location."),
    ("viewer_audit", "User and time of each location view/export.", "Admin-only; retain for accountability."),
]:
    row = table.add_row().cells
    for cell, value in zip(row, values): cell.text = value
table_anchor._p.addnext(table._tbl)

# Extend the existing comparison table, when present.
if doc.tables:
    comparison = doc.tables[0]
    row = comparison.add_row().cells
    values = ["Google Maps-Based Authorised Location Tracking", "No", "No", "Limited/Plugin", "Yes – consent-based last-known location, map, history and alerts"]
    for cell, value in zip(row, values): cell.text = value

doc.save(OUTPUT)
print(OUTPUT)
